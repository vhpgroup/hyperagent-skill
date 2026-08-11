from __future__ import annotations

from pathlib import Path

import pytest

from hsmt_engine.config import Settings


@pytest.fixture
def settings(tmp_path: Path) -> Settings:
    root = Path(__file__).resolve().parents[1]
    data = tmp_path / "data"
    data.mkdir()
    return Settings(
        root=root,
        data_dir=data,
        database=data / "jobs.sqlite3",
        api_token="test-token",
        api_host="127.0.0.1",
        api_port=8787,
        max_upload_mb=10,
        exa_api_key=None,
        llm_api_key=None,
        llm_base_url="https://api.openai.com/v1",
        llm_model=None,
        research_results=2,
    )


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    from docx import Document

    path = tmp_path / "sample-hsmt.docx"
    document = Document()
    document.add_heading("Gói thầu thiết bị văn phòng", level=1)
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "STT"
    table.rows[0].cells[1].text = "Tên hàng hóa"
    table.rows[0].cells[2].text = "Thông số kỹ thuật"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "Máy in laser"
    table.rows[1].cells[2].text = "Tốc độ: tối thiểu 30 trang/phút\nĐộ phân giải: 1200 dpi"
    document.save(path)
    return path
