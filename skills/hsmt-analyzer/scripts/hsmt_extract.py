#!/usr/bin/env python3
"""HSMT Analyzer — GĐ1: Bóc tách HSMT Phần 2 (DOCX) → JSON schema v7.

Usage:
  python3 hsmt_extract.py input.docx -o output.json

Yêu cầu: pip install python-docx
Output: JSON gồm project_info, general_requirements, items (components →
requirements với field/operator/value/unit/critical/weight/confidence),
quantity_table, validation_warnings. In thống kê ra stdout.

Hỗ trợ 2 layout bảng thông số (tự nhận diện):
  A — 3 cột kinh điển "STT | Tên hàng hóa | Thông số kỹ thuật" + bảng khối lượng riêng.
  B — gộp "STT | Hạng mục | ĐVT | SL": dòng đầu cell Hạng mục là tên, các dòng sau là
      thông số; SL lấy trực tiếp (spec_table), đối chiếu chéo bảng quy mô nếu có.
Tự LOẠI bảng mẫu tuyên bố đáp ứng (hàng "(1)(2)(3)…" / cell toàn "…").
Tách dòng thông số gộp bằng soft-separator " - Nhãn:" thành từng requirement riêng.
"""
import argparse
import json
import re


def norm(s):
    s = (s or "").lower().replace("\xa0", " ")
    return re.sub(r"\s+", " ", s).strip()


UNIT_RE = (r"(kva|kwh|kw|va|w|vac|vdc|v|ghz|mhz|hz|gbps|mbps|kbps|gb|tb|mb|"
           r"inch|mm|cm|km|m|kg|g|ah|ma|a|%|°c|dpi|ppm|fo|dbi|dbm|db|mp|nits|"
           r"cd/m2|u|lít|giờ|phút|giây|năm|tháng|ngày|cổng|port|core|nhân|luồng|pixel|px)")

FIELD_MAP = [
    (r"dạng sóng|sóng sine", "waveform"),
    (r"hệ số công suất|pf\s*[:=]?\s*[01]", "power_factor"),
    (r"công suất", "power_capacity"),
    (r"điện áp|voltage|vac|vdc", "voltage"),
    (r"tần số|frequency", "frequency"),
    (r"cổng kết nối|cổng|port|giao tiếp|kết nối|interface|rs232|usb|hdmi|vga|rj45|sc/upc|lc/upc", "interface"),
    (r"kích thước|dimension|(rộng|cao|sâu|dài)\b", "dimensions"),
    (r"trọng lượng|weight", "weight"),
    (r"độ phân giải|resolution", "resolution"),
    (r"màn hình|display|cảm ứng|inch", "display"),
    (r"cpu|chíp|chip|vi xử lý|processor|\bcore\b|nhân\b", "processor"),
    (r"\bram\b|bộ nhớ trong|memory", "memory"),
    (r"ổ cứng|ssd|hdd|lưu trữ|storage|emmc|rom", "storage"),
    (r"pin\b|battery|ắc quy|ăcquy|ăc quy|acquy", "battery"),
    (r"tốc độ|speed|scan.*ppm|trang/phút", "speed"),
    (r"tiêu chuẩn|chuẩn\b|standard|iso|iec|tcvn|\bul\b|\bce\b|ansi|astm", "standard"),
    (r"bảo hành|warranty", "warranty"),
    (r"xuất xứ|origin|hãng sản xuất|thương hiệu|brand", "origin"),
    (r"nhiệt độ|temperature|độ ẩm|humidity|môi trường", "environment"),
    (r"công nghệ|technology|\bonline\b|\boffline\b|line interactive", "product_type"),
    (r"wifi|wi-fi|băng tần|ăng ten|angten|anten|antenna|802\.11|\bax\b|\bac\b", "wireless"),
    (r"vlan|poe|qos|acl|snmp|layer|switching|throughput|bảng mac|\bmac\b|định tuyến|routing|igmp|stp", "network_feature"),
    (r"vật liệu|chất liệu|material|thép|nhựa|inox|tôn", "material"),
    (r"dung lượng", "capacity"),
    (r"thiết kế|dạng rack|khuôn dạng|form factor|kiểu dáng", "form_factor"),
]

CATS = [
    # Pattern đặc thù đặt TRƯỚC pattern rộng để ưu tiên khớp
    (r"máy in mã vạch", "barcode_printer", ["barcode label printer"], "high"),
    (r"máy đọc mã vạch|đầu đọc mã vạch", "barcode_scanner", ["barcode scanner"], "medium"),
    (r"chứng thư số|chữ ký số", "digital_certificate", ["digital certificate PKI"], "medium"),
    (r"kiosk", "kiosk", ["self-service kiosk"], "high"),
    (r"smart tivi|màn hình tivi|\btivi\b|smart tv", "tv_display", ["smart TV display"], "high"),
    (r"định tuyến|router", "router", ["network router"], "high"),
    (r"chấm công|kiểm soát ra vào", "access_control", ["access control time attendance"], "high"),
    (r"giám sát.{0,12}nhiệt độ|cảnh báo nhiệt độ", "env_monitoring", ["temperature humidity monitoring alarm"], "high"),
    (r"cửa.{0,12}chống cháy", "fire_door", ["fire rated steel door"], "low"),
    (r"sàn nâng", "raised_floor", ["raised access floor server room"], "medium"),
    (r"lưu điện|ups", "ups", ["UPS", "online UPS", "uninterruptible power supply"], "high"),
    (r"bóng chữa cháy|bình cầu", "fire_ball", ["fire extinguisher ball", "automatic fire suppression ball"], "high"),
    (r"báo cháy", "fire_alarm", ["fire alarm system", "smoke detector", "fire alarm control panel"], "high"),
    (r"cắt lọc sét|chống sét", "surge_protection", ["surge protection device", "lightning surge filter SPD"], "high"),
    (r"chuyển mạch|switch", "network_switch", ["L2 managed switch", "ethernet switch"], "high"),
    (r"máy chủ|server\b", "server", ["rack server"], "high"),
    (r"firewall|tường lửa", "firewall", ["network firewall"], "high"),
    (r"tủ máy chủ|tủ mạng|tủ rack", "rack_cabinet", ["server rack cabinet", "network rack"], "medium"),
    (r"dây cáp mạng|cat6|cat5", "network_cable", ["CAT6 UTP cable"], "low"),
    (r"máng ghen|máng cáp", "cable_trunking", ["PVC cable trunking"], "low"),
    (r"ổ cắm mạng", "network_outlet", ["RJ45 wall outlet faceplate"], "low"),
    (r"hạt mạng|rj45", "rj45_connector", ["RJ45 connector plug CAT6"], "low"),
    (r"ổ cắm điện", "power_outlet", ["power socket outlet"], "low"),
    (r"hộp otb", "fiber_box", ["optical termination box OTB"], "low"),
    (r"odf", "fiber_odf", ["ODF rack mount fiber distribution frame"], "low"),
    (r"cáp quang|\d+fo\b", "fiber_cable", ["fiber optic cable single mode"], "medium"),
    (r"sfp|module quang", "sfp_module", ["SFP transceiver module"], "medium"),
    (r"dây nhảy quang", "fiber_patchcord", ["fiber optic patch cord"], "low"),
    (r"wifi|access point", "wifi_ap", ["WiFi access point", "wireless access point"], "medium"),
    (r"nhân công|vật tư phụ", "labor_materials", ["installation labor materials"], "low"),
    (r"màn hình hiển thị|bảng thông báo|màn hình hỗ trợ", "display_signage", ["LCD display digital signage"], "medium"),
    (r"máy scan|máy quét", "scanner", ["document scanner ADF duplex"], "medium"),
    (r"xe tiêm", "medical_cart", ["medical laptop cart nursing trolley"], "low"),
    (r"máy tính bảng", "tablet", ["Android tablet"], "medium"),
    (r"ký điện tử|vân tay", "esign_fingerprint", ["signature pad", "fingerprint scanner"], "low"),
    (r"bộ máy tính|máy tính để bàn|máy vi tính", "desktop_pc", ["desktop computer PC"], "medium"),
    (r"laptop|máy tính xách tay", "laptop", ["business laptop"], "medium"),
    (r"diệt virus|antivirus", "antivirus", ["antivirus software license"], "low"),
    (r"máy in", "printer", ["laser printer"], "medium"),
    (r"camera", "camera", ["IP camera CCTV"], "medium"),
]


def detect_field(t):
    tl = norm(t)
    for pat, f in FIELD_MAP:
        if re.search(pat, tl):
            return f
    return "other"


def detect_operator(t):
    tl = norm(t)
    if re.search(r"≥|>=|tối thiểu|ít nhất|trở lên|không nhỏ hơn|lớn hơn hoặc bằng", tl):
        return ">="
    if re.search(r"≤|<=|tối đa|không quá|không lớn hơn|nhỏ hơn hoặc bằng", tl):
        return "<="
    if re.search(r"từ\s+[\d.,]+\s*(?:đến|~|tới)\s*[\d.,]+|[\d.,]+\s*~\s*[\d.,]+", tl):
        return "range"
    if re.search(r"hỗ trợ|tương thích|support", tl):
        return "supports"
    if re.search(r"\d+\s*/\s*\d+\s*/\s*\d+", tl):
        return "supports"
    if re.search(r"bao gồm|gồm|kèm theo|đi kèm", tl):
        return "contains_all"
    if detect_field(t) == "interface" and re.search(r"/|,| và ", tl):
        return "contains_all"
    return "="


def extract_value_unit(t):
    t2 = t.replace("\xa0", " ")
    m = re.search(r"(\d+(?:[.,]\d+)?)\s*" + UNIT_RE + r"\b", t2, re.I)
    if m:
        return m.group(1), m.group(2)
    m2 = re.search(r"\d+(?:[.,]\d+)?", t2)
    if m2:
        return m2.group(0), ""
    return None, ""


def detect_cat(name):
    nl = norm(name)
    for pat, cat, kws, prio in CATS:
        if re.search(pat, nl):
            return cat, kws, prio
    return "other", [], "medium"


TEMPLATE_NUM_RE = re.compile(r"^\(\d+\)$")


def is_template_table(rows):
    """Bảng MẪU tuyên bố đáp ứng (Chương V thường kèm) — KHÔNG phải bảng thông số thật.
    Đặc trưng: (a) có hàng đánh số cột '(1) (2) (3)…' ngay dưới header,
    hoặc (b) ≥40% cell thân bảng chỉ chứa dấu ba chấm '…'."""
    if not rows:
        return False
    for r in rows[:3]:
        cells = [c.strip() for c in r["dedup"] if c.strip()]
        if len(cells) >= 3 and sum(1 for c in cells if TEMPLATE_NUM_RE.match(c)) >= max(2, int(len(cells) * 0.6)):
            return True
    body = [c.strip() for r in rows[1:] for c in r["cells"] if c.strip()]
    if body:
        dots = sum(1 for c in body if set(c) <= set("…. …"))
        if dots / len(body) >= 0.4:
            return True
    return False


def looks_layout_b(rows):
    """Layout B: bảng 'STT | Hạng mục/Danh mục | ĐVT | SL' — thông số gộp trong cell cột 2
    (dòng đầu = tên hạng mục, các dòng sau = thông số), không có cột 'Thông số' riêng."""
    if not rows:
        return False
    head = " ".join(rows[0]["dedup"]).lower()
    return (("hạng mục" in head or "danh mục" in head)
            and ("đvt" in head or "đơn vị" in head)
            and ("sl" in head or "số lượng" in head)
            and "thông số" not in head)


def split_soft(ln):
    """Tách 1 dòng gộp nhiều thông số ngăn bằng ' - ' (soft separator) trước chữ hoa.
    VD: 'Công nghệ: Online - Hệ số CS: ≥99% - Điện áp vào: 220VAC' → 3 dòng."""
    points = [m for m in re.finditer(r"\s+-\s+", ln)
              if ln[m.end():m.end() + 1].isupper()]
    if len(points) < 2 or ln.count(":") < 2:
        return [ln]
    parts, last = [], 0
    for m in points:
        parts.append(ln[last:m.start()])
        last = m.end()
    parts.append(ln[last:])
    return [p.strip() for p in parts if p.strip()]


def split_spec_lines(text):
    """Tách khối thông số thành từng dòng: theo newline trước, rồi tách tiếp dòng gộp soft-separator."""
    out = []
    for ln in text.split("\n"):
        ln = ln.strip()
        if ln:
            out.extend(split_soft(ln))
    return out


def parse_qty_vn(s):
    """Parse số lượng kiểu VN: '1.900' (chấm ngàn) → 1900; '12,5' → 12.5."""
    t = (s or "").strip().replace(" ", "")
    if not t:
        return None
    if re.fullmatch(r"\d{1,3}(\.\d{3})+", t):
        t = t.replace(".", "")
    t = t.replace(",", ".")
    try:
        q = float(t)
        return int(q) if q == int(q) else q
    except ValueError:
        return None


def extract_docx(path):
    from docx import Document
    d = Document(path)
    paras = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    tables = []
    for t in d.tables:
        rows = []
        for r in t.rows:
            cells = [c.text for c in r.cells]
            dedup = []
            for c in cells:
                cc = " ".join(c.split())
                if not dedup or cc != dedup[-1]:
                    dedup.append(cc)
            rows.append({"cells": [c.text for c in r.cells], "dedup": dedup})
        tables.append(rows)
    # Loại bảng MẪU tuyên bố đáp ứng trước khi chọn (hàng '(1)(2)(3)…' / cell toàn '…')
    non_template = [rows for rows in tables if rows and not is_template_table(rows)]
    # Layout A: bảng thông số có header chứa "thông số"; bảng khối lượng chứa "khối lượng"/"mã hiệu"
    spec_t, qty_t, layout, qty_b = None, None, "A", None
    for rows in non_template:
        head = " ".join(rows[0]["dedup"]).lower()
        if spec_t is None and "thông số" in head:
            spec_t = rows
        elif qty_t is None and ("khối lượng" in head or "mã hiệu" in head):
            qty_t = rows
    # Layout B: 'STT | Hạng mục | ĐVT | SL' — bảng có cell mô tả dài nhất là bảng thông số,
    # bảng B còn lại (nếu có) là bảng quy mô dùng đối chiếu số lượng
    if spec_t is None:
        b_tables = [rows for rows in non_template if looks_layout_b(rows)]
        if b_tables:
            def desc_len(rows):
                vals = [len(r["cells"][1]) for r in rows[1:] if len(r["cells"]) > 1]
                return sum(vals) / len(vals) if vals else 0
            spec_t = max(b_tables, key=desc_len)
            layout = "B"
            rest = [b for b in b_tables if b is not spec_t]
            if rest:
                qty_b = max(rest, key=len)
    if spec_t is None and non_template:
        spec_t = max(non_template, key=lambda rs: max((len(r["cells"][-1]) for r in rs), default=0))
    return paras, spec_t or [], qty_t or [], layout, qty_b or []


def parse(docx_path):
    paras, spec_rows, qty_rows, layout, qty_b_rows = extract_docx(docx_path)

    raw_items = []
    if layout == "B":
        # Layout B: cột 2 = tên (dòng đầu) + thông số (các dòng sau); cột 3/4 = ĐVT/SL
        for row in spec_rows[1:]:
            cells = row["cells"]
            if len(cells) < 4:
                continue
            lines_b = [l.strip() for l in cells[1].split("\n") if l.strip()]
            if not lines_b:
                continue
            name_b = " ".join(lines_b[0].rstrip(":").split())
            if not name_b:
                continue
            raw_items.append({"stt": cells[0].strip(), "name": name_b,
                              "spec_raw": "\n".join(lines_b[1:]),
                              "unit_b": " ".join(cells[2].split()),
                              "qty_b": " ".join(cells[3].split())})
    else:
        for row in spec_rows[1:]:
            cells = row["cells"]
            if len(cells) < 3:
                continue
            stt = cells[0].strip()
            name = " ".join(cells[1].split())
            spec = cells[2]
            if not (stt or name or spec.strip()):
                continue
            raw_items.append({"stt": stt, "name": name, "spec_raw": spec})

    # quantity table + breakdowns
    qt_entries, bd_list, cur = [], [], None
    for row in qty_rows:
        r = row["dedup"]
        first = r[0].strip() if r else ""
        if re.fullmatch(r"\d+", first):
            cur = {"stt": int(first), "code": r[1] if len(r) > 1 else "",
                   "description": r[2] if len(r) > 2 else "",
                   "unit": r[3] if len(r) > 3 else "",
                   "quantity": r[4] if len(r) > 4 else "", "breakdown": []}
            qt_entries.append(cur)
        else:
            text = " ".join(c for c in r if c).strip()
            if cur is not None and text and not text.upper().startswith("STT"):
                cur["breakdown"].append(text)
                m = re.match(r"(.+?)\s*:\s*([\d.,\sx*+=/()\-]+)$", text)
                if m:
                    total = re.findall(r"=\s*([\d.,]+)", m.group(2))
                    bd_list.append({"key": norm(m.group(1)), "qty": total[-1] if total else None,
                                    "unit": cur["unit"], "used": False, "main_qty": cur["quantity"]})

    # Layout B: bảng quy mô 'STT | Danh mục | ĐVT | SL' làm quantity_table đối chiếu chéo
    if layout == "B" and qty_b_rows:
        for row in qty_b_rows[1:]:
            c = row["cells"]
            if len(c) >= 4 and c[0].strip():
                qt_entries.append({"stt": c[0].strip(), "code": "",
                                   "description": " ".join(c[1].split()),
                                   "unit": " ".join(c[2].split()),
                                   "quantity": " ".join(c[3].split()), "breakdown": []})

    def find_val(prefix):
        for p in paras:
            pn = norm(p)
            if pn.startswith("- " + prefix) or pn.startswith("-" + prefix):
                return p.split(":", 1)[1].strip() if ":" in p else p
        return ""

    project_info = {
        "investor": find_val("chủ đầu tư"), "address": find_val("địa chỉ"),
        "package_name": find_val("tên gói thầu"), "funding_source": find_val("nguồn vốn"),
        "procurement_method": find_val("phương thức lựa chọn nhà thầu"),
        "contract_type": find_val("loại hợp đồng"), "location": find_val("địa điểm thực hiện"),
        "implementation_time": find_val("thời gian thực hiện gói thầu") or find_val("thời gian thực hiện"),
    }

    general_requirements, in_gen = [], False
    for i, p in enumerate(paras):
        pn = norm(p)
        if re.match(r"^1\.2", p) or pn.startswith("yêu cầu về kỹ thuật chung"):
            in_gen = True
            continue
        if pn.startswith("mục 2"):
            break
        if pn.startswith("ghi chú") or pn.startswith("yêu cầu về kỹ thuật cụ thể"):
            in_gen = False
            continue
        if in_gen and (p.startswith("-") or (i > 0 and paras[i - 1].startswith("-")
                       and not re.match(r"^(\d|[a-z]\)|mục)", pn))):
            general_requirements.append(p.lstrip("- ").strip())

    items, warnings, name_seen = [], [], {}
    for idx, ri in enumerate(raw_items, start=1):
        name = ri["name"]
        cat, en_kws, prio = detect_cat(name)
        lines = split_spec_lines(ri["spec_raw"])

        has_star_hdr = any(re.match(r"^\*\s+\S", l) for l in lines)
        n_plus = sum(1 for l in lines if re.match(r"^\+\s", l))
        dash_lines = [l for l in lines if re.match(r"^-\s", l)]

        def dash_no_colon_content(l):
            body = re.sub(r"^-\s*", "", l)
            if ":" not in body:
                return True
            return body.split(":", 1)[1].strip() == ""

        dash_hdr_cand = [l for l in dash_lines if dash_no_colon_content(l)]
        dash_is_header = (len(dash_hdr_cand) >= 2 and n_plus >= 2
                          and len(dash_hdr_cand) > len(dash_lines) / 2)

        def header_of(line):
            mh = re.match(r"^(\d+)\s*/\s*(.+)", line)
            if mh:
                return mh.group(2).strip()
            if has_star_hdr:
                ms = re.match(r"^\*\s+(.+)", line)
                if ms:
                    return ms.group(1).strip()
            if dash_is_header and re.match(r"^-\s", line) and dash_no_colon_content(line):
                return re.sub(r"^-\s*", "", line).strip().rstrip(":")
            return None

        comps, cur_comp = [], None
        for line in lines:
            h = header_of(line)
            if h:
                cur_comp = {"header": h, "lines": []}
                comps.append(cur_comp)
            else:
                if cur_comp is None:
                    cur_comp = {"header": None, "lines": []}
                    comps.append(cur_comp)
                cur_comp["lines"].append(line)
        if not comps:
            comps = [{"header": None, "lines": lines}]

        components = []
        for k, c in enumerate(comps, start=1):
            cname = c["header"] or name
            reqs = []
            for line in c["lines"]:
                clean = re.sub(r"^[.\-+*•]\s*", "", line).strip()
                if not clean:
                    continue
                if reqs and ":" not in clean and re.match(r"^[a-zà-ỹđ0-9]", clean):
                    reqs[-1]["raw_text"] += " " + line
                    continue
                if clean.endswith(":") and len(clean) <= 45 and not re.search(r"\d", clean):
                    reqs.append({"raw_text": line, "field": "section_label", "operator": "=",
                                 "value": None, "unit": "", "required": False,
                                 "critical": False, "weight": 2, "confidence": 0.5})
                    continue
                field = detect_field(clean)
                op = detect_operator(clean)
                val, unit = extract_value_unit(clean)
                has_num = val is not None
                conf = 0.9 if (has_num and field != "other") else (0.6 if has_num or field != "other" else 0.4)
                critical = bool(has_num and field in {
                    "power_capacity", "power_factor", "voltage", "frequency", "interface",
                    "resolution", "display", "processor", "memory", "storage", "speed",
                    "capacity", "wireless", "network_feature", "battery", "product_type"})
                if field in {"standard", "product_type", "form_factor"} and not has_num:
                    critical = True
                weight = 10 if (critical and has_num) else (8 if critical else (3 if field in {"warranty", "origin", "environment", "material"} else 5))
                if conf <= 0.4:
                    warnings.append({"type": "requirement_confidence_thap", "item_no": idx,
                                     "component_no": f"{idx}.{k}", "raw_text": line,
                                     "issue": "Không nhận diện được field/giá trị rõ ràng",
                                     "suggested_fix": "Rà tay khi so khớp"})
                reqs.append({"raw_text": line, "field": field, "operator": op,
                             "value": val, "unit": unit, "required": True,
                             "critical": critical, "weight": weight, "confidence": conf})
            if not reqs and c["header"]:
                hf = detect_field(c["header"])
                hv, hu = extract_value_unit(c["header"])
                reqs.append({"raw_text": c["header"], "field": hf, "operator": "=",
                             "value": hv, "unit": hu, "required": True,
                             "critical": False, "weight": 5, "confidence": 0.6})
            key_specs = [r["value"] + r["unit"] for r in reqs if r["critical"] and r["value"]][:3]
            components.append({
                "component_no": f"{idx}.{k}", "name": cname,
                "component_type": "service" if cat in {"labor_materials", "digital_certificate"} else "product",
                "category": cat,
                "search_keywords": list(dict.fromkeys([cname, name] + en_kws)),
                "search_query": " ".join([name] + en_kws[:1] + key_specs),
                "equivalent_allowed": True,
                "evidence_required": ["catalogue", "datasheet"] + (["certificate"] if cat in {"fire_ball", "fire_alarm", "surge_protection"} else []),
                "requirements": reqs,
            })

        qty, unit, qsrc, qnote = None, "", "", ""
        nn = norm(name)
        if layout == "B":
            # Số lượng nằm ngay trên bảng thông số (cột ĐVT/SL); đối chiếu chéo bảng quy mô nếu có
            qty = parse_qty_vn(ri.get("qty_b"))
            unit, qsrc = ri.get("unit_b", ""), "spec_table"
            ref = [q for q in qt_entries
                   if norm(q["description"])[:40] == nn[:40]
                   or nn and nn[:40] in norm(q["description"]) or norm(q["description"])[:40] in nn]
            if ref and norm(str(ref[0]["quantity"])) != norm(str(ri.get("qty_b", ""))):
                qnote = f"Lệch SL với bảng quy mô: {ref[0]['quantity']} vs {ri.get('qty_b')}"
                warnings.append({"type": "quantity_chua_chac", "item_no": idx, "component_no": "",
                                 "raw_text": name, "issue": qnote,
                                 "suggested_fix": "Đối chiếu 2 bảng trong HSMT"})
        else:
            cand = [b for b in bd_list if not b["used"] and (b["key"] == nn or nn in b["key"] or b["key"] in nn)]
            if cand:
                b = cand[0]
                b["used"] = True
                try:
                    q = float(b["qty"]) if b["qty"] else float(b["main_qty"])
                    qty = int(q) if q == int(q) else q
                except (TypeError, ValueError):
                    qty = None
                unit, qsrc = b["unit"], "quantity_table"
            else:
                # Fallback: match trực tiếp dòng chính bảng khối lượng theo mô tả
                ref = [q for q in qt_entries if q.get("description")
                       and (norm(q["description"]) == nn or nn in norm(q["description"])
                            or norm(q["description"]) in nn)]
                qv = parse_qty_vn(str(ref[0]["quantity"])) if ref else None
                if qv is not None:
                    qty, unit, qsrc = qv, ref[0]["unit"], "quantity_table"
                else:
                    warnings.append({"type": "quantity_khong_tim_thay", "item_no": idx, "component_no": "",
                                     "raw_text": name, "issue": "Không tìm thấy dòng khối lượng khớp tên",
                                     "suggested_fix": "Đối chiếu tay với bảng khối lượng"})
                    qnote = "Chưa khớp được với bảng khối lượng"

        if nn in name_seen:
            warnings.append({"type": "item_trung", "item_no": idx, "component_no": "",
                             "raw_text": name, "issue": f"Trùng tên với hạng mục {name_seen[nn]}",
                             "suggested_fix": "Kiểm tra vị trí/mục đích khác nhau"})
        name_seen.setdefault(nn, idx)

        crit_texts = [r["raw_text"] for c in components for r in c["requirements"] if r["critical"]][:10]
        items.append({
            "item_no": idx, "item_name": name,
            "item_type": "service" if cat in {"labor_materials", "digital_certificate"} else "goods",
            "category": cat, "quantity": qty, "unit": unit,
            "quantity_source": qsrc, "quantity_note": qnote,
            "match_priority": prio, "must_have_summary": crit_texts,
            "raw_block": ri["spec_raw"], "components": components,
        })

    by_name = {}
    for it in items:
        by_name.setdefault(norm(it["item_name"]), []).append(it)
    for group in by_name.values():
        if len(group) > 1:
            with_qty = [g for g in group if g["quantity"] is not None]
            without = [g for g in group if g["quantity"] is None]
            if len(with_qty) == 1 and without:
                src = with_qty[0]
                note = (f"Bảng khối lượng chỉ có 1 dòng gộp chung ({src['quantity']} {src['unit']}) "
                        f"cho {len(group)} hạng mục trùng tên (số {', '.join(str(g['item_no']) for g in group)})")
                for g in group:
                    g["quantity_note"] = (g["quantity_note"] + "; " if g["quantity_note"] else "") + note
                warnings.append({"type": "quantity_gop_chung", "item_no": without[0]["item_no"],
                                 "component_no": "", "raw_text": src["item_name"],
                                 "issue": note, "suggested_fix": "Phân bổ theo thiết kế"})

    return {"project_info": project_info, "general_requirements": general_requirements,
            "items": items, "quantity_table": qt_entries, "validation_warnings": warnings}


def main():
    ap = argparse.ArgumentParser(description="Bóc tách HSMT DOCX -> JSON v7")
    ap.add_argument("docx")
    ap.add_argument("-o", "--output", default="hsmt_extraction.json")
    a = ap.parse_args()
    out = parse(a.docx)
    json.dump(out, open(a.output, "w", encoding="utf-8"), ensure_ascii=False, indent=1)
    n_req = sum(len(c["requirements"]) for it in out["items"] for c in it["components"])
    n_comp = sum(len(it["components"]) for it in out["items"])
    n_crit = sum(1 for it in out["items"] for c in it["components"] for r in c["requirements"] if r["critical"])
    print(f"OK {a.output}: {len(out['items'])} items | {n_comp} components | "
          f"{n_req} requirements ({n_crit} critical) | {len(out['validation_warnings'])} warnings | "
          f"{len(out['general_requirements'])} general reqs | {len(out['quantity_table'])} quantity rows")


if __name__ == "__main__":
    main()
